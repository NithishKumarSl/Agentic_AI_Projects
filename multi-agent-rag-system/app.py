import os
import streamlit as st
import pdfplumber
from docx import Document as DocxDocument

from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langgraph.graph import StateGraph
from langchain_core.runnables import RunnableLambda

from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.tools import DuckDuckGoSearchRun
from dotenv import load_dotenv
load_dotenv()
# ---------------------- CONFIGURATION ----------------------
GOOGLE_API_KEY = os.getenv("GEMINI_API_KEY")

if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in environment variables")

llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.2, google_api_key=GOOGLE_API_KEY)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
search = DuckDuckGoSearchRun()

# ---------------------- FILE PARSER ----------------------
def extract_text_from_local_path(path):
    if path.endswith(".pdf"):
        with pdfplumber.open(path) as pdf:
            return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    elif path.endswith(".txt"):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    elif path.endswith(".docx"):
        doc = DocxDocument(path)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

# ---------------------- VECTOR STORE / RETRIEVER ----------------------
def build_retriever():
    docs = []
    if os.path.exists("my_docs"):
        for fname in os.listdir("my_docs"):
            fpath = os.path.join("my_docs", fname)
            if fname.lower().endswith((".pdf", ".txt", ".docx")):
                text = extract_text_from_local_path(fpath)
                if text:
                    for chunk in text_splitter.split_text(text):
                        docs.append(Document(page_content=chunk, metadata={"source": fname}))
    if docs:
        db = FAISS.from_documents(docs, embeddings)
        return db.as_retriever()
    return None

retriever = build_retriever()

# ---------------------- AGENTS ----------------------
def router_agent(state):
    query = state.get("query", "")
    route_prompt = PromptTemplate.from_template(
        "Classify the query into one of [web, rag, llm]:\n\nQuery: {query}\n\nAnswer:"
    )
    route_result = (route_prompt | llm).invoke({"query": query}).content.lower()
    route = "llm"
    if "web" in route_result:
        route = "web"
    elif "rag" in route_result:
        route = "rag"
    return {**state, "route": route}

def web_agent(state):
    query = state["query"]
    try:
        result = search.run(query)
        return {**state, "content": result}
    except Exception as e:
        return {**state, "content": f"Web search failed: {str(e)}"}

def rag_agent(state):
    query = state["query"]
    retriever = state["retriever"]
    qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)
    answer = qa_chain.run(query)
    return {**state, "content": answer}

def llm_agent(state):
    query = state["query"]
    response = llm.invoke(query)
    return {**state, "content": response.content}

def summarizer_agent(state):
    content = state["content"]
    prompt = PromptTemplate.from_template("Summarize clearly and concisely:\n\n{content}")
    summary = (prompt | llm).invoke({"content": content}).content
    return {**state, "final": summary}

# ---------------------- LANGGRAPH ----------------------
def run_langgraph(user_query, retriever):
    workflow = StateGraph(dict)
    workflow.set_entry_point("router")

    workflow.add_node("router", RunnableLambda(router_agent))
    workflow.add_node("web", RunnableLambda(web_agent))
    workflow.add_node("rag", RunnableLambda(rag_agent))
    workflow.add_node("llm", RunnableLambda(llm_agent))
    workflow.add_node("summarizer", RunnableLambda(summarizer_agent))

    def router_logic(state): return state["route"]
    workflow.add_conditional_edges("router", router_logic, {
        "web": "web",
        "rag": "rag",
        "llm": "llm"
    })

    for node in ["web", "rag", "llm"]:
        workflow.add_edge(node, "summarizer")

    workflow.set_finish_point("summarizer")
    app = workflow.compile()
    return app.invoke({"query": user_query, "retriever": retriever})["final"]

# ---------------------- STREAMLIT APP ----------------------
# --- Custom CSS for modern look ---
st.markdown(
    """
    <style>
    .main {
        background-color: #f7f9fa;
        border-radius: 18px;
        padding: 2rem 2rem 1rem 2rem;
        box-shadow: 0 2px 16px rgba(0,0,0,0.07);
    }
    .stButton>button {
        background: linear-gradient(90deg, #4f8cff 0%, #6edb8f 100%);
        color: white;
        border-radius: 8px;
        font-weight: bold;
        font-size: 1.1rem;
        padding: 0.5rem 2rem;
        margin-top: 1rem;
    }
    .stTextInput>div>div>input {
        border-radius: 8px;
        border: 1.5px solid #4f8cff;
        font-size: 1.1rem;
        padding: 0.5rem 1rem;
    }
    .stAlert {
        border-radius: 8px;
    }
    footer {
        visibility: hidden;
    }
    .custom-footer {
        position: fixed;
        left: 0; right: 0; bottom: 0;
        width: 100%;
        background: #e9ecef;
        color: #333;
        text-align: center;
        padding: 0.5rem 0;
        font-size: 0.95rem;
        z-index: 100;
        border-top: 1px solid #d1d5db;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# --- Sidebar ---
st.sidebar.title("‚ÑπÔ∏è About")
st.sidebar.markdown(
    """
    **Multi-Agent RAG System**
    
    - Combines Web Search, RAG, and LLM
    - Powered by LangGraph, Gemini, and FAISS
    - Upload your own documents to the `my_docs` folder
    """
)

st.sidebar.markdown("---")

# Document status in sidebar
if os.path.exists("my_docs"):
    doc_files = [f for f in os.listdir("my_docs") if f.lower().endswith((".pdf", ".txt", ".docx"))]
    st.sidebar.success(f"üìÇ {len(doc_files)} document(s) found in 'my_docs'.")
else:
    st.sidebar.info("No 'my_docs' folder found. Using fallback knowledge base.")

st.sidebar.markdown("---")
st.sidebar.markdown("Created by [Your Name] ¬∑ Powered by Streamlit & LangChain")

# --- Main UI ---
st.markdown("""
<div class="main">
    <h1 style="text-align:center; font-size:2.5rem; margin-bottom:0.2em;">üß† Multi-Agent RAG System</h1>
    <p style="text-align:center; color:#4f8cff; font-size:1.2rem; margin-bottom:2em;">
        <b>LangGraph + Web + RAG + LLM</b> ‚Äî Your fully agentic research assistant
    </p>
</div>
""", unsafe_allow_html=True)

# --- Columns for input/output ---
col1, col2 = st.columns([1,2])

with col1:
    st.markdown("#### üí¨ Ask your question")
    query = st.text_input("", placeholder="e.g. What is LangGraph?", key="user_query")
    submit = st.button("Submit")

with col2:
    if submit:
        if not query.strip():
            st.warning("‚ö†Ô∏è Please enter a question.")
        else:
            with st.spinner("ü§ñ Thinking..."):
                try:
                    answer = run_langgraph(query, retriever)
                    st.success("‚úÖ Done!")
                    st.subheader("üìò Answer:")
                    st.write(answer)
                except Exception as e:
                    st.error(f"‚ùå Error: {str(e)}")
    else:
        st.info("Enter a question and click Submit to get started.")

# --- Footer ---
st.markdown(
    """
    <div class="custom-footer">
        &copy; 2024 Multi-Agent RAG System &mdash; <a href="https://streamlit.io/" target="_blank">Streamlit</a> + <a href="https://python.langchain.com/" target="_blank">LangChain</a> + <a href="https://ai.google.dev/gemini-api/docs" target="_blank">Gemini</a>
    </div>
    """,
    unsafe_allow_html=True,
) 
